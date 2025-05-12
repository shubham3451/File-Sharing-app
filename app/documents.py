import fitz  
import docx  
from django_elasticsearch_dsl import Document, Index, fields
from django_elasticsearch_dsl.registries import registry
from models.file import FileVersion


file_versions_index = Index('fileversions')

@registry.register_document
class FileVersionDocument(Document):
    # Index setup
    class Index:
        name = 'fileversions'
    
    class Django:
        model = FileVersion 
        fields = ['file__name', 'file__created_at', 'uploaded_at']

    def get_file_text(self):
        """Dynamically extract text based on file type"""
        file_path = self.instance.file_data.path 
        text = ''

        if file_path.endswith('.pdf'):
            text = self.extract_pdf_text(file_path)
        elif file_path.endswith('.docx'):
            text = self.extract_docx_text(file_path)
        elif file_path.endswith('.txt'):
            text = self.extract_text_file(file_path)
        # Add more types (like pptx, xls, etc.) if needed

        return text

    def extract_pdf_text(self, file_path):
        """Extract text from PDF using PyMuPDF."""
        try:
            doc = fitz.open(file_path)
            text = ''
            for page in doc:
                text += page.get_text()
            return text
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""

    def extract_docx_text(self, file_path):
        """Extract text from DOCX file using python-docx."""
        try:
            doc = docx.Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""

    def extract_text_file(self, file_path):
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from file: {e}")
            return ""
